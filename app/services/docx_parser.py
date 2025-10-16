from typing import Dict, Any, List
from io import BytesIO
from docx import Document
import re

from app.utils.regex_patterns import FinancialPatterns
from app.utils.text_utils import TextProcessor


class DocxParser:
    """Rule-based parser for structured DOCX documents"""

    def __init__(self):
        self.patterns = FinancialPatterns()
        self.text_processor = TextProcessor()

    def parse(self, content: bytes) -> Dict[str, Any]:
        """Parse DOCX document and extract financial entities"""
        doc = Document(BytesIO(content))

        full_text = self._extract_text(doc)
        table_data = self._extract_from_tables(doc)
        pattern_entities = self.patterns.extract_all(full_text)
        kv_entities = self._extract_key_value_pairs(full_text)

        return self._merge_entities(pattern_entities, table_data, kv_entities)

    def _extract_text(self, doc: Document) -> str:
        """Extract all text from document paragraphs"""
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = '\n'.join(paragraphs)
        return self.text_processor.clean_text(full_text)

    def _extract_from_tables(self, doc: Document) -> Dict[str, List[str]]:
        """Extract entities from document tables"""
        entities = {
            'counterparty': [],
            'notional': [],
            'isin': [],
            'underlying': [],
            'maturity': [],
            'coupon': [],
            'barrier': [],
            'trade_date': [],
            'currency': []
        }

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                for i in range(len(cells) - 1):
                    key = cells[i].lower()
                    value = cells[i + 1]

                    if any(term in key for term in ['counterparty', 'party', 'issuer']):
                        if value and len(value) > 2:
                            entities['counterparty'].append(value)

                    elif any(term in key for term in ['notional', 'principal', 'amount']):
                        numeric = re.findall(r'[\d,]+\.?\d*', value)
                        if numeric:
                            entities['notional'].extend(numeric)

                    elif 'isin' in key:
                        isin_match = self.patterns.ISIN.findall(value)
                        entities['isin'].extend(isin_match)

                    elif any(term in key for term in ['underlying', 'reference', 'asset']):
                        if value and len(value) > 2:
                            entities['underlying'].append(value)

                    elif any(term in key for term in ['maturity', 'expiry', 'expiration']):
                        date_match = self.patterns.DATE.findall(value)
                        if date_match:
                            entities['maturity'].extend(date_match)
                        elif value:
                            entities['maturity'].append(value)

                    elif any(term in key for term in ['coupon', 'rate', 'interest']):
                        percent_match = self.patterns.PERCENTAGE.findall(value)
                        if percent_match:
                            entities['coupon'].extend(percent_match)

                    elif any(term in key for term in ['barrier', 'strike', 'trigger']):
                        numeric = re.findall(r'[\d,]+\.?\d*', value)
                        if numeric:
                            entities['barrier'].extend(numeric)

                    elif any(term in key for term in ['trade date', 'execution date']):
                        date_match = self.patterns.DATE.findall(value)
                        entities['trade_date'].extend(date_match)

                    elif 'currency' in key:
                        if value and len(value) == 3 and value.isupper():
                            entities['currency'].append(value)

        return entities

    def _extract_key_value_pairs(self, text: str) -> Dict[str, List[str]]:
        """Extract entities from key-value patterns (Key: Value or Key - Value)"""
        entities = {
            'counterparty': [],
            'notional': [],
            'underlying': [],
            'maturity': [],
            'coupon': [],
            'barrier': [],
            'trade_date': [],
            'payment_frequency': []
        }

        lines = text.split('\n')

        for line in lines:
            line = line.strip()

            if ':' in line or '-' in line:
                parts = re.split(r'[:\-]', line, 1)
                if len(parts) == 2:
                    key = parts[0].strip().lower()
                    value = parts[1].strip()

                    if any(term in key for term in ['counterparty', 'party', 'issuer', 'client']):
                        entities['counterparty'].append(value)

                    elif any(term in key for term in ['notional', 'principal', 'nominal']):
                        entities['notional'].append(value)

                    elif any(term in key for term in ['underlying', 'reference', 'asset']):
                        entities['underlying'].append(value)

                    elif any(term in key for term in ['maturity', 'expiry', 'expiration']):
                        entities['maturity'].append(value)

                    elif any(term in key for term in ['coupon', 'interest rate']):
                        entities['coupon'].append(value)

                    elif any(term in key for term in ['barrier', 'strike', 'trigger']):
                        entities['barrier'].append(value)

                    elif any(term in key for term in ['trade date', 'execution']):
                        entities['trade_date'].append(value)

                    elif any(term in key for term in ['payment frequency', 'frequency']):
                        entities['payment_frequency'].append(value)

        return entities

    def _merge_entities(self, *entity_dicts) -> Dict[str, List[str]]:
        """Merge and deduplicate entities from multiple sources"""
        merged = {}
        all_types = set()

        for d in entity_dicts:
            all_types.update(d.keys())

        for entity_type in all_types:
            values = []
            seen = set()

            for entity_dict in entity_dicts:
                if entity_type in entity_dict:
                    items = entity_dict[entity_type]
                    if not isinstance(items, list):
                        items = [items]

                    for item in items:
                        normalized = self.text_processor.normalize_entity(str(item))
                        if normalized and normalized.lower() not in seen:
                            values.append(normalized)
                            seen.add(normalized.lower())

            if values:
                merged[entity_type] = values

        return merged
